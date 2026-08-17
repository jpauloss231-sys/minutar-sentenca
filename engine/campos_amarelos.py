#!/usr/bin/env python3
"""
Biblioteca para localizar, extrair e preencher campos em destaque amarelo
dentro de modelos .docx reais (minutas oficiais do magistrado).

Um "campo amarelo" é um grupo de runs consecutivos com
WD_COLOR_INDEX.YELLOW dentro do mesmo parágrafo (ou célula de tabela).
Cada grupo recebe um índice sequencial (chunk) dentro do parágrafo/célula,
na ordem em que aparece.
"""
import json
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def _agrupar_runs_amarelos(paragraph):
    """Retorna lista de grupos: cada grupo é uma lista de índices de runs
    consecutivos com destaque amarelo dentro do parágrafo."""
    grupos = []
    atual = []
    for i, r in enumerate(paragraph.runs):
        if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
            atual.append(i)
        else:
            if atual:
                grupos.append(atual)
                atual = []
    if atual:
        grupos.append(atual)
    return grupos


def listar_campos(path):
    """Lista todos os campos amarelos do documento (parágrafos + tabelas)."""
    doc = Document(path)
    campos = []

    for pi, p in enumerate(doc.paragraphs):
        for ci, grupo in enumerate(_agrupar_runs_amarelos(p)):
            texto = "".join(p.runs[i].text for i in grupo)
            campos.append({
                "local": {"tipo": "paragrafo", "indice": pi},
                "chunk": ci,
                "texto_original": texto,
            })

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for cj, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    for ci, grupo in enumerate(_agrupar_runs_amarelos(p)):
                        texto = "".join(p.runs[i].text for i in grupo)
                        campos.append({
                            "local": {"tipo": "tabela", "tabela": ti, "linha": ri,
                                      "coluna": cj, "paragrafo": pi},
                            "chunk": ci,
                            "texto_original": texto,
                        })
    return campos


def _get_paragraph(doc, local):
    if local["tipo"] == "paragrafo":
        return doc.paragraphs[local["indice"]]
    else:
        table = doc.tables[local["tabela"]]
        cell = table.rows[local["linha"]].cells[local["coluna"]]
        return cell.paragraphs[local["paragrafo"]]


def preencher(template_path, mapeamento, valores, out_path, manter_destaque=True):
    """
    template_path: caminho do .docx original (será apenas lido)
    mapeamento: lista de dicts {"local": {...}, "chunk": N, "campo": "chave"}
                (mesmo formato retornado por listar_campos, mas com "campo"
                indicando a chave em `valores` que deve substituir o texto)
    valores: dict {"chave": "texto novo"}
    out_path: caminho de saída do .docx preenchido
    manter_destaque: se False, remove o destaque amarelo após preencher
                      (útil para gerar a versão "limpa" pronta para assinar)
    """
    doc = Document(template_path)

    # Agrupa o mapeamento por (local, chunk) para achar os runs certos
    for item in mapeamento:
        campo = item["campo"]
        if campo not in valores:
            continue  # mantém o texto original do modelo se não houver valor
        novo_texto = valores[campo]
        p = _get_paragraph(doc, item["local"])
        grupos = _agrupar_runs_amarelos(p)
        if item["chunk"] >= len(grupos):
            continue
        grupo = grupos[item["chunk"]]
        # Escreve o novo texto no primeiro run do grupo, limpa os demais
        primeiro = True
        for i in grupo:
            run = p.runs[i]
            if primeiro:
                run.text = novo_texto
                if not manter_destaque:
                    run.font.highlight_color = None
                primeiro = False
            else:
                run.text = ""

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    import sys
    campos = listar_campos(sys.argv[1])
    print(json.dumps(campos, ensure_ascii=False, indent=2))
