#!/usr/bin/env python3
"""
render.py - Motor de renderização de minutas de sentença em lote.

Uso:
    python3 render.py --data dados.json --out ../output

Onde dados.json é uma lista de casos, cada um com:
    {
        "id": "processo_0001234-12.2024.4.01.8000",
        "template": "auxilio_doenca_invalidez",   # ou "bpc_loas"
        "fields": { ... campos exigidos pelo template ... }
    }

Os modelos ficam em ../templates/*.txt (sintaxe Jinja2) e o
registro de modelos disponíveis em ../templates/registry.json.
Para adicionar um novo modelo de sentença, basta criar um novo
arquivo .txt em templates/ e cadastrá-lo no registry.json.
"""
import argparse
import json
import os
import sys
import zipfile
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"


def load_registry():
    with open(TEMPLATES_DIR / "registry.json", encoding="utf-8") as f:
        return json.load(f)["templates"]


def render_text(template_id: str, fields: dict) -> str:
    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    registry = {t["id"]: t for t in load_registry()}
    if template_id not in registry:
        raise ValueError(f"Modelo desconhecido: {template_id}")
    tpl = env.get_template(registry[template_id]["arquivo"])
    return tpl.render(**fields)


def text_to_docx(text: str, out_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if line else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(10)
        if line:
            p.paragraph_format.first_line_indent = Cm(0) if line.isupper() else Cm(1.25)
            run = p.add_run(line)
            if line.isupper() and len(line) < 60:
                run.bold = True
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, help="JSON com a lista de casos")
    parser.add_argument("--out", default=str(BASE_DIR / "output"), help="Pasta de saída")
    parser.add_argument("--zip", action="store_true", help="Gerar também um .zip com todas as minutas")
    args = parser.parse_args()

    with open(args.data, encoding="utf-8") as f:
        cases = json.load(f)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    generated = []
    for case in cases:
        case_id = case["id"]
        template_id = case["template"]
        fields = case["fields"]
        text = render_text(template_id, fields)
        out_path = out_dir / f"{case_id}.docx"
        text_to_docx(text, out_path)
        generated.append(out_path)
        print(f"[ok] {case_id} -> {out_path.name} (modelo: {template_id})")

    if args.zip and generated:
        zip_path = out_dir / "minutas.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for p in generated:
                zf.write(p, arcname=p.name)
        print(f"[ok] pacote gerado: {zip_path}")

    print(f"\nTotal de minutas geradas: {len(generated)}")


if __name__ == "__main__":
    main()
